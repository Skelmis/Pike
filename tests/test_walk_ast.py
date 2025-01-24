from pathlib import Path
from unittest.mock import Mock, call

import pytest
from docx import Document
from docx.text.run import Run

from pike import Engine, File, utils, jinja_globals
from pike.docx import Docx


def test_nested_blockquote_doesnt_crash(engine: Engine, data_dir) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "nested_blockquote.md").read_text())
    docx.walk_ast(Document(), ast)


def test_line_break(engine: Engine, data_dir) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "line_break.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_run("Line break:"),
        call.add_paragraph().add_run(),
        call.add_paragraph().add_run().add_break(),
        call.add_paragraph().add_run("This is the first line."),
        call.add_paragraph().add_run(),
        call.add_paragraph().add_run().add_break(),
        call.add_paragraph().add_run("And this is the second line."),
    ]


def test_paragraphs(engine: Engine, data_dir) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "paragraphs.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_run("I really like using Markdown."),
        call.add_paragraph(),
        call.add_paragraph().add_run(
            "I think I'll use it to format all of my documents from now on."
        ),
    ]


def test_multiple_ordered_lists(engine: Engine, data_dir) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "ordered_lists.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().restart_numbering(),
        call.add_paragraph().add_run("First item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Second item"),
        call.add_paragraph(),
        call.add_paragraph().add_run("Next list:"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().restart_numbering(),
        call.add_paragraph().add_run("First item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Second item"),
    ]


def test_ordered_list_nested_restarting(engine: Engine, data_dir) -> None:
    """Odd edge case, restarted the first two ordered lists fine but not nested"""
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "multi_ordered_list.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_run("Ordered lists:"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().restart_numbering(),
        call.add_paragraph().add_run("First item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Second item"),
        call.add_paragraph(),
        call.add_paragraph().add_run("Different markdown, same list result:"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().restart_numbering(),
        call.add_paragraph().add_run("First item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Second item"),
        call.add_paragraph(),
        call.add_paragraph().add_run("Unordered lists:"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("First item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Second item"),
        call.add_paragraph(),
        call.add_paragraph().add_run("Nested lists:"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().restart_numbering(),
        call.add_paragraph().add_run("First item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Second item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Third item"),
        call.add_paragraph(style="List Paragraph 2"),
        call.add_paragraph().add_run("Indented item"),
        call.add_paragraph(style="List Paragraph 2"),
        call.add_paragraph().add_run("Indented item"),
        call.add_paragraph(style="List Paragraph"),
        call.add_paragraph().add_run("Fourth item"),
    ]


def test_inline_code(engine: Engine, data_dir) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "inline_code.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_run(style="Subtle Reference"),
        call.add_paragraph().add_run().add_text("code"),
    ]


def test_horizontal_rule(engine: Engine, data_dir) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "horizontal_rule.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().draw_paragraph_border(top=True),
        call.add_paragraph(),
        call.add_paragraph().draw_paragraph_border(top=True),
        call.add_paragraph(),
        call.add_paragraph().draw_paragraph_border(top=True),
    ]


def test_quick_links(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "quick_link.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_external_hyperlink(
            "https://google.com", "https://google.com"
        ),
    ]


def test_normal_links(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "normal_link.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_external_hyperlink("https://google.com", "Google"),
    ]


def test_titled_link(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "titled_link.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_external_hyperlink("https://google.com", "Google"),
    ]


def test_image(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "image.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_picture("/images/cat.jpg", width=None, height=None),
    ]


@pytest.mark.xfail(
    reason="Until we implement caption and alt text this isnt finished",
    strict=True,
)
def test_image_with_sizes(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "sized_image.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [call.add_picture("cat.jpg", width=2, height=2)]
    # TODO Finish this
    raise ValueError("unfinished")


def test_image_via_jinja_inline_insert_image(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    img_str = jinja_globals.insert_image(
        "images/cat.jpg",
        width=1,
        height=2,
    )
    ast = markdown.parse(f"Hello{img_str}")
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_paragraph(),
        call.add_paragraph().add_run("Hello"),
        call.add_picture("images/cat.jpg", width=360000, height=720000),
    ]


def test_image_via_jinja_block_insert_image(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse(
        jinja_globals.insert_image(
            "images/cat.jpg",
            width=1,
            height=2,
        )
    )
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_picture("images/cat.jpg", width=360000, height=720000)
    ]


def test_bold_title(engine: Engine, data_dir: Path) -> None:
    # This is kind of a fake test? At time of writing I couldn't
    # figure out how to mock down to attributes being set so
    # this was manually tested and then we just assume if
    # the mock calls are made then this is #fine
    docx = Docx(engine)
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "bold_title.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [
        call.add_heading(level=1),
        call.add_heading().add_run("Bold"),
        call.add_heading().add_run(" nothing "),
        call.add_heading().add_run("italic"),
    ]


def test_custom_command(engine: Engine, data_dir: Path) -> None:
    docx = Docx(engine)
    docx.import_commands_from_engine()
    markdown = utils.create_markdown_it()
    ast = markdown.parse((data_dir / "custom_command.md").read_text())
    document = Mock()
    docx.walk_ast(document, ast)
    assert document.mock_calls == [call.add_page_break()]
